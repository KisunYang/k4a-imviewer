# -*- coding: utf-8 -*-
"""개발보고서.md -> 개발보고서.docx 변환 (python-docx)"""
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Inches
except ImportError:
    sys.exit("pip install python-docx 필요")


def add_inline_runs(paragraph, text: str) -> None:
    """**굵게** 와 `코드` 처리."""
    # **` 같이 섞이면 단순 순회
    rest = text
    while rest:
        m = re.search(r"(\*\*[^*]+\*\*|`[^`]+`)", rest)
        if not m:
            paragraph.add_run(rest)
            break
        if m.start() > 0:
            paragraph.add_run(rest[: m.start()])
        token = m.group(1)
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        else:
            r = paragraph.add_run(token[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        rest = rest[m.end() :]


def md_to_docx(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    sect = doc.sections[0]
    sect.page_height = Inches(11.69)
    sect.page_width = Inches(8.27)
    sect.left_margin = Inches(1)
    sect.right_margin = Inches(1)

    i = 0
    in_code = False
    code_lines = []

    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        # 표: | 로 시작하는 연속 줄
        if line.strip().startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = lines[i].strip()
                if re.match(r"^\|[\s\-:|]+\|$", row.replace(" ", "")):
                    i += 1
                    continue
                cells = [c.strip() for c in row.strip("|").split("|")]
                rows.append(cells)
                i += 1
            if rows:
                tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
                tbl.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        if ci < len(tbl.rows[ri].cells):
                            tbl.rows[ri].cells[ci].text = cell
            continue

        if line.startswith("### "):
            h = doc.add_heading(line[4:].strip(), level=2)
            h.style.font.name = "맑은 고딕"
            i += 1
            continue
        if line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=1)
            h.style.font.name = "맑은 고딕"
            i += 1
            continue
        if line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h.style.font.name = "맑은 고딕"
            i += 1
            continue

        if re.match(r"^\d+\.\s", line.strip()):
            text = re.sub(r"^\d+\.\s*", "", line.strip())
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, text)
            i += 1
            continue

        if line.strip().startswith("- "):
            text = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, text)
            i += 1
            continue

        if line.strip().startswith("*") and line.strip().endswith("*") and line.strip().count("*") == 2:
            inner = line.strip()[1:-1]
            p = doc.add_paragraph()
            r = p.add_run(inner)
            r.italic = True
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline_runs(p, line.strip())
        i += 1

    doc.save(out_path)


def main():
    root = Path(__file__).resolve().parent.parent
    md = root / "개발보고서.md"
    out = root / "개발보고서.docx"
    if not md.exists():
        sys.exit(f"없음: {md}")
    md_to_docx(md, out)
    print(out)


if __name__ == "__main__":
    main()
